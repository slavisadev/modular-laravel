import sys
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

# Add title slide
title = doc.add_heading('Laravel Packages Scenarios', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('A comprehensive overview of Laravel package development')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Loop through all the slide files
presentation_dir = os.path.dirname(os.path.abspath(__file__))

for i in range(1, 11):
    slide_path = os.path.join(presentation_dir, f'slide{i}.md')
    
    if os.path.exists(slide_path):
        with open(slide_path, 'r') as f:
            content = f.read()
        
        # Split content by lines
        lines = content.strip().split('\n')
        
        # First line is the heading
        doc.add_page_break()
        heading = doc.add_heading(lines[0].replace('# ', ''), level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        current_bullet = None
        code_block = False
        code_content = []
        
        # Process the remaining lines
        for line in lines[2:]:  # Skip first line (title) and empty second line
            # Check if this is the start of a code block
            if line.strip() == '```' or line.strip().startswith('```'):
                code_block = not code_block
                if not code_block and code_content:  # End of code block
                    code_text = '\n'.join(code_content)
                    p = doc.add_paragraph()
                    code_run = p.add_run(code_text)
                    code_run.font.name = 'Courier New'
                    code_run.font.size = Pt(9)
                    code_content = []
                continue
            
            # If inside code block, collect the content
            if code_block:
                code_content.append(line)
                continue
            
            # Check if this is a bullet point
            if line.strip().startswith('- **'):
                # Extract the bold text
                bold_text = line.strip().replace('- **', '').split('**')[0]
                current_bullet = doc.add_paragraph(style='List Bullet')
                current_bullet.add_run(bold_text).bold = True
            
            # Check if this is a sub-bullet under the current bullet
            elif line.strip().startswith('  -') and current_bullet:
                sub_bullet_text = line.strip().replace('  - ', '')
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.5)
                p.add_run(sub_bullet_text)

# Save the document
doc.save(os.path.join(presentation_dir, 'Laravel_Packages_Presentation.docx'))
print("Document created successfully!")